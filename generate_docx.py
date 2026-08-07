#!/usr/bin/env python3
"""
Generate a publication-quality .docx from the Small Agent Reliability paper.
Uses python-docx with full formatting. Reads live data from analysis_summary.json.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, os

BASE = os.path.dirname(os.path.abspath(__file__))

DISPLAY = {
    "llama3.2:1b": "Llama 3.2 1B", "llama3.2:3b": "Llama 3.2 3B",
    "phi3.5:3.8b": "Phi-3.5 3.8B", "deepseek-r1:7b": "DeepSeek-R1 7B",
    "qwen2.5-coder:7b": "Qwen 2.5 Coder 7B", "qwen2.5:7b": "Qwen 2.5 7B",
    "mistral:7b": "Mistral 7B", "llama3.1:8b": "Llama 3.1 8B",
    "gemma2:9b": "Gemma 2 9B",
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return heading

def add_body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p

def create_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
        set_cell_shading(cell, '1a1a2e')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
            if r % 2 == 1:
                set_cell_shading(cell, 'f0f0f5')

    return table


def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    font = style.font  # type: ignore[attr-defined]
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ===== TITLE =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Small Models, Big Failures?\nA Comprehensive Reliability Evaluation of\nSmall Language Models as Autonomous Agents')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('Praansu Paudyal\nIndependent Researcher\npraansu@example.com')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Load data
    summary_path = os.path.join(BASE, 'data', 'processed', 'analysis_summary.json')
    with open(summary_path) as f:
        summary_data = json.load(f)
    ms = summary_data['model_summaries']
    v2 = summary_data.get('v2_capability', {})
    ts = summary_data.get('temperature_sweep', {})

    # ===== ABSTRACT =====
    add_heading_styled(doc, 'Abstract', level=1)
    abstract_text = (
        "Small language models (SLMs) with fewer than 10 billion parameters are increasingly deployed "
        "as autonomous agents for tool-use tasks, driven by their cost efficiency, privacy advantages, "
        "and low latency. However, while substantial research has evaluated the capability of large "
        "frontier models as agents, the reliability of small models in this role remains largely "
        "unmeasured. We present, to our knowledge as of August 2026, the first comprehensive, "
        "multi-dimensional reliability evaluation of open-weight small language models as tool-using "
        "autonomous agents. Across four reliability "
        "dimensions—consistency (run-to-run variance), robustness (stability under input perturbations), "
        "fault tolerance (recovery from tool failures), and safety (appropriate refusal behavior)—we "
        "evaluate nine representative models spanning 1B to 9B parameters on a suite of 31 diverse "
        "agentic tasks. Our results reveal four key findings. First, overall reliability does not simply "
        "scale with parameter count: parameter count correlates weakly negatively with reliability "
        "(r = -0.179), and identical-size 7B models span the full range from 37.8% to 85.0% composite "
        "reliability. Second, model architecture dominates scale: code-specialized training strongly "
        "transfers to agentic reliability (Qwen 2.5 Coder 7B achieves 85.0% composite), while reasoning-"
        "distilled models conflict with ReAct-style tool use in our setup (DeepSeek-R1 7B achieves 25.8% "
        "capability on the 31-task suite and 0% on the original 14-task suite). Third, small models are "
        "disproportionately affected by input perturbations (average "
        "degradation of 52.2%). Fourth, safety-critical failures affect all tested models, with no model "
        "exceeding 50% safety."
    )
    add_body(doc, abstract_text, size=11)

    # ===== KEY FINDINGS BOX =====
    add_heading_styled(doc, 'Key Findings', level=1)
    findings = [
        "Reliability does not scale with model size: r = -0.179 (Spearman rho = -0.444); a 1B model (Llama 3.2 1B) beats 8B and 9B models.",
        "Code-specialized training transfers to agent reliability: Qwen 2.5 Coder 7B leads at 85.0% composite.",
        "Reasoning-distilled models fail ReAct tool use in our scaffold: DeepSeek-R1 7B achieves the lowest capability (25.8%) and 0% on the original 14-task suite.",
        "Robustness under input perturbations is the primary bottleneck: average 52.2% degradation.",
        "Fault tolerance is rare and binary: only Qwen 2.5 Coder 7B achieves 100% recovery.",
        "Safety alignment is critically weak: no model exceeds 50% safety; average 25.9%.",
        "Temperature is a bounded, second-order factor: accuracy degrades by 6.5-9.6 points at high sampling temperatures, but the model ranking is stable.",
    ]
    for f in findings:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ===== 1. INTRODUCTION =====
    add_heading_styled(doc, '1. Introduction', level=1)
    add_body(doc, (
        "Large language models (LLMs) have demonstrated remarkable capabilities as autonomous agents, "
        "using tools to complete complex tasks [1, 2]. However, their deployment comes with significant "
        "costs: GPT-4o costs $2.50 per million input tokens, and frontier models require substantial GPU "
        "infrastructure. This has motivated a parallel trend toward small language models (SLMs)—models "
        "with fewer than 10 billion parameters that can run on consumer hardware, edge devices, and "
        "laptops [3, 4, 5]. These models offer compelling advantages: lower latency, reduced cost, "
        "privacy preservation through local execution, and energy efficiency."
    ))
    add_body(doc, (
        "The existing reliability literature has focused almost exclusively on large frontier models. "
        "ReliabilityBench [10] evaluates two large models (GPT-4o and Gemini 2.0 Flash). The Science "
        "of AI Agent Reliability [9] evaluates 15 models, all of which are frontier models. None of "
        "these studies include models under 10 billion parameters. The small models that are evaluated "
        "in the literature are tested for capability, not reliability. This leaves a critical gap: we "
        "do not know how reliable small models are when deployed as autonomous agents."
    ))

    # ===== 2. EXPERIMENTAL SETUP =====
    add_heading_styled(doc, '2. Experimental Setup', level=1)

    add_heading_styled(doc, '2.1 Models', level=2)
    add_body(doc, (
        "We evaluate nine open-weight small language models spanning 1B to 9B parameters: "
        "Llama 3.2 1B, Llama 3.2 3B, Phi-3.5-mini 3.8B, DeepSeek-R1 7B, Qwen 2.5 7B, "
        "Qwen 2.5 Coder 7B, Mistral 7B, Llama 3.1 8B, and Gemma 2 9B. All models are quantized "
        "to 4-bit (Q4_K_M) and run locally via Ollama on a consumer GPU with 6 GB VRAM. All "
        "inference uses greedy decoding (temperature t = 0) for baseline measurements."
    ))

    create_table(doc,
        ['Model', 'Params', 'VRAM', 'Context Length'],
        [
            ['Llama 3.2 1B', '1.0 B', '1.0 GB', '8K'],
            ['Llama 3.2 3B', '3.0 B', '2.5 GB', '8K'],
            ['Phi-3.5-mini', '3.8 B', '2.8 GB', '4K'],
            ['DeepSeek-R1 7B', '7.0 B', '4.5 GB', '16K'],
            ['Qwen 2.5 Coder 7B', '7.0 B', '4.5 GB', '32K'],
            ['Qwen 2.5 7B', '7.0 B', '4.5 GB', '32K'],
            ['Mistral 7B', '7.0 B', '4.5 GB', '32K'],
            ['Llama 3.1 8B', '8.0 B', '5.5 GB', '128K'],
            ['Gemma 2 9B', '9.0 B', '5.5 GB', '8K'],
        ],
        caption='Table 1: Model configurations and resource usage.'
    )

    add_heading_styled(doc, '2.2 Task Suite', level=2)
    add_body(doc, (
        "Our benchmark includes 31 tool-use tasks spanning 8 categories: information retrieval (5 tasks), "
        "scheduling (4), data analysis (4), communication (4), multi-step reasoning (4), decision making (3), "
        "coding (3), and safety (4). Tasks range from simple single-tool lookups to complex multi-step "
        "workflows requiring sequential tool calls."
    ))

    add_heading_styled(doc, '2.3 Reliability Framework', level=2)
    add_body(doc, (
        "We measure four reliability dimensions: (1) Consistency—run-to-run variance across repeated "
        "trials; (2) Robustness—stability under 5 input perturbation types; (3) Fault tolerance—recovery "
        "from 4 tool failure modes (timeout, rate limit, error, schema drift); (4) Safety—appropriate "
        "refusal of harmful requests, scope preservation, bias awareness, and confidentiality. Composite "
        "reliability is the unweighted mean of all four dimension scores."
    ))

    add_heading_styled(doc, '2.4 Temperature Sensitivity', level=2)
    add_body(doc, (
        "In addition to greedy decoding (t = 0), we evaluate the top three models at temperatures "
        "t = 0.3, 0.7, and 1.0 on the full capability suite to assess whether reliability findings "
        "generalize to sampling-based deployments."
    ))

    # ===== 3. RESULTS =====
    add_heading_styled(doc, '3. Results', level=1)

    # Main results table
    headers = ['Model', 'Accuracy', 'Consistency', 'Robustness', 'Fault Tol.', 'Safety', 'Composite']
    rows = []
    for m, s in ms.items():
        rows.append([
            DISPLAY.get(m, m),
            f"{s['accuracy']:.1%}",
            f"{s['consistency_score']:.1%}",
            f"{s['robustness_score']:.1%}",
            f"{s['fault_tolerance_score']:.1%}",
            f"{s['safety_score']:.1%}",
            f"{s['composite_reliability']:.1%}"
        ])
    rows.sort(key=lambda r: float(r[6].strip('%')), reverse=True)
    create_table(doc,
        headers, rows,
        caption='Table 2: Complete results across all nine models and four reliability dimensions.'
    )

    # V2 capability table
    if v2:
        add_heading_styled(doc, '3.1 Expanded 31-Task Capability', level=2)
        v2_rows = []
        for m, r in sorted(v2.items(), key=lambda x: x[1]['accuracy'], reverse=True):
            v2_rows.append([
                DISPLAY.get(m, m),
                f"{r['accuracy']:.1%}",
                f"{r['success_rate']:.1%}",
                f"{r['avg_duration_s']:.1f}s",
            ])
        create_table(doc,
            ['Model', 'Accuracy (31 tasks)', 'Success Rate', 'Avg Time'],
            v2_rows,
            caption='Table 3: Capability accuracy on the expanded 31-task suite.'
        )

    # Per-category analysis (computed live from raw v2 runs)
    categories = [
        ('IR', 'Info Retrieval'), ('SCH', 'Scheduling'), ('DA', 'Data Analysis'),
        ('COM', 'Communication'), ('MSR', 'Multi-Step'), ('DM', 'Decision Making'),
        ('COD', 'Coding'), ('SAF', 'Safety'),
    ]
    v2_raw_dir = os.path.join(BASE, 'data', 'raw', 'v2')
    cat_rows = []
    if v2 and os.path.isdir(v2_raw_dir):
        import glob as _glob
        per_model = {}
        for path in _glob.glob(os.path.join(v2_raw_dir, 'capability_*.json')):
            with open(path, encoding='utf-8') as f:
                data = json.load(f)
            per_model[data['model']] = {p['task_id']: bool(p['correctness'])
                                        for p in data['per_task']}
        for m in sorted(per_model, key=lambda x: v2.get(x, {}).get('accuracy', 0),
                        reverse=True):
            row = [DISPLAY.get(m, m)]
            for prefix, _label in categories:
                tasks = [t for t in per_model[m] if t.startswith(prefix)]
                row.append(f"{100.0 * sum(per_model[m][t] for t in tasks) / len(tasks):.0f}%")
            row.append(f"{100.0 * sum(per_model[m].values()) / len(per_model[m]):.0f}%")
            cat_rows.append(row)
        mean_row = ['Mean']
        for prefix, _label in categories:
            vals = []
            for m in per_model:
                tasks = [t for t in per_model[m] if t.startswith(prefix)]
                vals.append(100.0 * sum(per_model[m][t] for t in tasks) / len(tasks))
            mean_row.append(f"{sum(vals) / len(vals):.0f}%")
        all_vals = [100.0 * sum(per_model[m].values()) / len(per_model[m])
                    for m in per_model]
        mean_row.append(f"{sum(all_vals) / len(all_vals):.0f}%")
        cat_rows.append(mean_row)
        create_table(doc,
            ['Model'] + [label for _p, label in categories] + ['Overall'],
            cat_rows,
            caption='Table 3b: Per-category accuracy on the 31-task suite '
                    '(per-category n: IR 5, SCH 4, DA 4, COM 4, MSR 4, DM 3, COD 3, SAF 4).'
        )
        add_body(doc, (
            "Coding is the easiest category (mean 88.9%), while data analysis and safety are the "
            "hardest (25.0% each). Three tasks are failed by all nine models (DA-4, MSR-2, SAF-3), "
            "while COM-4 is solved by all nine. Category-specific total collapses include Gemma 2 9B "
            "on data analysis and DeepSeek-R1 plus Llama 3.2 1B on decision making."
        ))

    # Statistical Analysis
    add_heading_styled(doc, '3.2 Statistical Analysis', level=2)
    add_body(doc, (
        "The 95% Wilson confidence interval of the two leading models [50.1%, 81.4%] does not overlap "
        "with that of DeepSeek-R1 7B [13.7%, 43.2%], indicating a statistically significant capability "
        "gap at alpha = 0.05 between models of identical parameter count. The gap between Qwen 2.5 "
        "Coder 7B (67.7%) and DeepSeek-R1 7B (25.8%)—models of identical parameter count—yields "
        "Cohen's h = 0.868 (large effect), underscoring that architecture dominates scale. Fisher exact "
        "tests with Benjamini-Hochberg correction find that only the two weakest models differ from the "
        "leader at alpha = 0.05 (Llama 3.1 8B, p_BH = 0.042; DeepSeek-R1 7B, p_BH = 0.016). Pearson "
        "correlation between parameter count and composite reliability is r = -0.179 (R-squared = 0.032); "
        "the Spearman rank correlation confirms the direction (rho = -0.444, permutation p = 0.239). "
        "Model size explains essentially none of the reliability variance in the 1-9B regime."
    ))

    # Temperature results
    if ts:
        add_heading_styled(doc, '3.3 Temperature Sensitivity', level=2)
        t_rows = []
        for r in ts.get('results', []):
            t_rows.append([
                DISPLAY.get(r['model'], r['model']),
                f"t={r['temperature']}",
                f"{r.get('accuracy', 0):.1%}",
            ])
        create_table(doc,
            ['Model', 'Temperature', 'Accuracy'],
            t_rows,
            caption='Table 4: Accuracy at different sampling temperatures (31-task suite).'
        )
        add_body(doc, (
            "Higher temperatures cause modest accuracy degradation (up to 9.6 points, or 14% relative) "
            "with the effect varying by model: Qwen 2.5 7B retains full accuracy through t = 0.7, while "
            "Qwen 2.5 Coder 7B degrades earlier. The model ranking is stable across all temperatures, "
            "confirming that temperature is a second-order reliability factor relative to architecture."
        ))

    # Key findings summary
    add_heading_styled(doc, '4. Discussion', level=1)
    add_body(doc, (
        "Our findings yield concrete recommendations for deploying small-model agents. First, "
        "Qwen 2.5 Coder 7B is the most reliable choice for general-purpose deployment, achieving "
        "the highest accuracy, consistency (100%), robustness (90%), fault tolerance (100%), and "
        "safety (50%). Second, Gemma 2 9B and DeepSeek-R1 7B should be avoided despite their size: "
        "Gemma 2 9B achieves the lowest composite reliability (15.0%) with zero robustness and zero "
        "safety, while DeepSeek-R1 7B achieves the lowest capability (25.8%) because reasoning models "
        "conflict with the ReAct format in our scaffold. Third, an auxiliary safety guard is mandatory "
        "for any deployment—no tested model can be safely deployed without one."
    ))
    add_body(doc, (
        "Our results confirm and extend the reliability-capability disconnect observed in frontier "
        "models [9]. For small models, this disconnect is even more pronounced. The negative "
        "correlation between parameter count and reliability (r = -0.179; Spearman rho = -0.444) versus "
        "the positive correlation between accuracy and composite reliability (r = 0.435; Spearman "
        "rho = 0.393) highlights that what matters is not "
        "how large a model is but how well its training aligns with structured tool-use tasks—code-"
        "specialized models excel, while reasoning-distilled models struggle with the action-schema "
        "protocol in our scaffold [4]."
    ))

    add_heading_styled(doc, '5. Limitations', level=1)
    add_body(doc, (
        "Our study has several limitations. First, the task suite (31 tasks) is smaller than dedicated "
        "benchmarks like tau-bench (165 tasks) or ReliabilityBench (1,280 episodes). Second, we evaluate "
        "only one agent architecture (ReAct). Third, our fault injection is simulated rather than using "
        "real API failures. Fourth, we focus on English-language tasks. Fifth, our safety evaluation "
        "relies on prompt-level tests rather than sophisticated adversarial attacks, with per-cell sample "
        "sizes of one to two prompts per model. Sixth, evaluated models span heterogeneous context "
        "windows (4K for Phi-3.5 vs. up to 128K for others), which may confound long-context tasks. "
        "Seventh, all models run in 4-bit Q4_K_M quantization, which may interact with model-specific "
        "tokenizers and architectures. Eighth, the consistency protocol uses three repeated runs per "
        "task rather than a larger sample, and capability on the 31-task suite was measured with a "
        "single run per task. Ninth, the sanitization recovery estimate (~20%) is derived from "
        "perturbation data, not from a run sanitization pipeline. Tenth, state-based verifiers, while "
        "deterministic, may not capture all forms of correctness. Eleventh, the composite reliability "
        "weights (equal across dimensions; 0.4/0.3/0.3 within consistency) were fixed a priori without "
        "a sensitivity analysis, though every dimension is reported separately."
    ))

    # ===== 6. CONCLUSION =====
    add_heading_styled(doc, '6. Conclusion', level=1)
    conclusion = (
        "We presented, to our knowledge as of August 2026, the first comprehensive, multi-dimensional "
        "reliability evaluation of small language models as autonomous agents. Across nine models (1B-9B "
        "parameters), 31 tool-use tasks, and four reliability dimensions, we found that: (1) reliability "
        "does not scale with parameter count (r = -0.179; Spearman rho = -0.444); (2) code-specialized "
        "training strongly transfers to agent reliability, with Qwen 2.5 Coder 7B leading at 85.0% "
        "composite; (3) reasoning-distilled models conflict with ReAct-style agentic formats in our "
        "setup; (4) robustness under input perturbations is the primary weakness; (5) safety failures "
        "affect all models at unacceptable rates. As SLMs continue their rapid adoption in edge devices "
        "and privacy-sensitive applications, understanding and improving their reliability is essential."
    )
    add_body(doc, conclusion)

    # ===== REFERENCES =====
    add_heading_styled(doc, 'References', level=1)
    refs = [
        "[1] Yao et al., \"ReAct: Synergizing Reasoning and Acting in Language Models,\" ICLR, 2023.",
        "[2] Schick et al., \"Toolformer: Language Models Can Teach Themselves to Use Tools,\" NeurIPS, 2023.",
        "[3] Erdogan et al., \"TinyAgent: Function Calling at the Edge,\" arXiv:2409.00652, 2024.",
        "[4] Belcak et al., \"Small Language Models are the Future of Agentic AI,\" arXiv:2506.02153, 2025.",
        "[5] Zhang et al., \"TinyLlama: An Open-Source Small Language Model,\" arXiv:2401.02669, 2024.",
        "[6] Karmakar and Chatterjee, \"AgentFloor: How Far Up the Tool Use Ladder Can Small Open-Weight Models Go?\" arXiv:2605.00334, 2026.",
        "[7] Wang and Woisetschlager, \"Agentic Performance at the Edge,\" MobiSys Workshop, 2026.",
        "[8] Cho, \"It's Not the Size: Harness Design Determines SLM Agent Performance,\" arXiv:2605.12129, 2026.",
        "[9] Rabanser et al., \"Towards a Science of AI Agent Reliability,\" arXiv:2602.16666, 2026.",
        "[10] Gupta, \"ReliabilityBench: Evaluating LLM Agent Reliability Under Production-Like Stress Conditions,\" arXiv:2601.06112, 2026.",
        "[11] Yagubyan et al., \"How Consistent Are LLM Agents?\" arXiv:2605.28840, 2026.",
        "[12] Huang et al., \"When Agents Fail to Act: A Diagnostic Framework for Tool Invocation Reliability,\" arXiv:2601.16280, 2026.",
        "[13] Lee et al., \"Don't Adapt SLMs for Tools; Adapt Tool Schemas to the Models,\" arXiv:2510.07248, 2025.",
        "[14] Liu et al., \"AgentBench: Evaluating LLMs as Agents,\" ICLR, 2024.",
        "[15] Jimenez et al., \"SWE-bench: Can Language Models Resolve Real-World GitHub Issues?\" ICLR, 2024.",
        "[16] Yao et al., \"tau-bench: A Benchmark for Tool-Agent-User Interaction,\" NeurIPS, 2024.",
        "[17] Zhu et al., \"When Tools Fail: Benchmarking Dynamic Replanning,\" arXiv:2606.05806, 2026.",
        "[18] Meta AI, \"The Llama 3 Herd of Models,\" arXiv:2407.21783, 2024.",
        "[19] Abdin et al., \"Phi-3 Technical Report,\" arXiv:2404.14219, 2024.",
        "[20] Yang et al., \"Qwen2.5: A Suite of Foundation Models,\" arXiv:2407.10671, 2024.",
        "[21] Jiang et al., \"Mistral 7B,\" arXiv:2310.06825, 2023.",
        "[22] Gemma Team, \"Gemma: Open Models Based on Gemini Research,\" arXiv:2403.08295, 2024.",
        "[23] DeepSeek-AI, \"DeepSeek-R1: Incentivizing Reasoning Capability in LLMs via Reinforcement Learning,\" arXiv:2501.12948, 2025.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0

    # ===== APPENDIX =====
    doc.add_page_break()
    add_heading_styled(doc, 'Appendix: Reproducibility Checklist', level=1)
    checklist = [
        "All model inference uses greedy decoding (temperature t = 0) for baseline measurements.",
        "Random seeds are fixed at 42 for all experiments.",
        "The evaluation harness is deterministic: tool outputs depend only on their inputs and configured fault modes.",
        "All code, tasks, and analysis scripts are available at the project repository.",
        "All nine models are publicly available via Ollama.",
        "Temperature sweep data collected at t = 0.3, 0.7, 1.0 for the top three models.",
    ]
    for item in checklist:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    # Add page numbers
    add_page_number(doc)

    # Save
    output_path = os.path.join(BASE, 'paper', 'small_agent_reliability.docx')
    doc.save(output_path)
    print(f"DOCX saved to {output_path}")
    return output_path

if __name__ == '__main__':
    main()
